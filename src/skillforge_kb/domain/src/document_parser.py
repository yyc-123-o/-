from __future__ import annotations

import hashlib
import os
import re
from dataclasses import dataclass, field
from typing import List, Optional

import fitz
from docx import Document as DocxDocument


@dataclass
class DocBlock:
    doc_id: str
    block_id: str
    heading_path: List[str]
    heading_level: int
    text: str
    page_no: Optional[int] = None
    domain_tag: Optional[str] = None
    difficulty: Optional[str] = None


@dataclass
class ParsedDocument:
    doc_id: str
    source_path: str
    doc_type: str
    title: str
    blocks: List[DocBlock] = field(default_factory=list)


def _make_doc_id(path: str) -> str:
    return hashlib.md5(path.encode("utf-8")).hexdigest()[:12]


def _clean_text(text: str) -> str:
    text = text.replace("\u3000", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"\n\d{1,3}\n", "\n", text)
    return text.strip()


def parse_pdf(path: str, domain_tag: str | None = None) -> ParsedDocument:
    doc_id = _make_doc_id(path)
    pdf = fitz.open(path)
    title = os.path.splitext(os.path.basename(path))[0]
    blocks: List[DocBlock] = []
    heading_stack: List[str] = []

    for page_no, page in enumerate(pdf, start=1):
        page_dict = page.get_text("dict")
        for block in page_dict.get("blocks", []):
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                line_text = "".join(span.get("text", "") for span in spans).strip()
                if not line_text:
                    continue

                max_size = max((span.get("size", 0) for span in spans), default=0)
                is_heading = max_size >= 13.5 and len(line_text) <= 48

                if is_heading:
                    level = 1 if max_size >= 16 else 2
                    heading_stack = heading_stack[: level - 1] + [line_text]
                    continue

                cleaned = _clean_text(line_text)
                if not cleaned:
                    continue

                if blocks and blocks[-1].heading_path == heading_stack and blocks[-1].page_no == page_no:
                    blocks[-1].text = f"{blocks[-1].text} {cleaned}".strip()
                else:
                    blocks.append(
                        DocBlock(
                            doc_id=doc_id,
                            block_id=f"{doc_id}-{page_no}-{len(blocks)}",
                            heading_path=list(heading_stack),
                            heading_level=0,
                            text=cleaned,
                            page_no=page_no,
                            domain_tag=domain_tag,
                        )
                    )

    return ParsedDocument(doc_id=doc_id, source_path=path, doc_type="pdf", title=title, blocks=blocks)


def parse_docx(path: str, domain_tag: str | None = None) -> ParsedDocument:
    doc_id = _make_doc_id(path)
    docx = DocxDocument(path)
    title = os.path.splitext(os.path.basename(path))[0]
    blocks: List[DocBlock] = []
    heading_stack: List[str] = []

    for para in docx.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = (para.style.name or "").lower()
        if "heading" in style_name:
            match = re.search(r"\d+", style_name)
            level = int(match.group()) if match else 1
            heading_stack = heading_stack[: level - 1] + [text]
            continue

        blocks.append(
            DocBlock(
                doc_id=doc_id,
                block_id=f"{doc_id}-{len(blocks)}",
                heading_path=list(heading_stack),
                heading_level=0,
                text=_clean_text(text),
                domain_tag=domain_tag,
            )
        )

    return ParsedDocument(doc_id=doc_id, source_path=path, doc_type="docx", title=title, blocks=blocks)


def parse_markdown(path: str, domain_tag: str | None = None) -> ParsedDocument:
    doc_id = _make_doc_id(path)
    title = os.path.splitext(os.path.basename(path))[0]
    blocks: List[DocBlock] = []
    heading_stack: List[str] = []
    current_lines: List[str] = []

    def flush() -> None:
        if not current_lines:
            return
        text = _clean_text("\n".join(current_lines))
        current_lines.clear()
        if not text:
            return
        blocks.append(
            DocBlock(
                doc_id=doc_id,
                block_id=f"{doc_id}-{len(blocks)}",
                heading_path=list(heading_stack),
                heading_level=0,
                text=text,
                domain_tag=domain_tag,
            )
        )

    with open(path, "r", encoding="utf-8") as fh:
        for line in fh:
            match = re.match(r"^(#{1,6})\s+(.*)$", line.rstrip("\n"))
            if match:
                flush()
                level = len(match.group(1))
                heading = match.group(2).strip()
                heading_stack = heading_stack[: level - 1] + [heading]
            else:
                current_lines.append(line.rstrip("\n"))
    flush()

    return ParsedDocument(doc_id=doc_id, source_path=path, doc_type="markdown", title=title, blocks=blocks)


def parse_document(path: str, domain_tag: str | None = None) -> ParsedDocument:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return parse_pdf(path, domain_tag)
    if ext == ".docx":
        return parse_docx(path, domain_tag)
    if ext in {".md", ".markdown"}:
        return parse_markdown(path, domain_tag)
    raise ValueError(f"Unsupported file type: {ext}")
